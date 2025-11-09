# src/export.py
from __future__ import annotations
from io import BytesIO
from datetime import datetime
from typing import Iterable
import pandas as pd

# DOCX
from docx import Document
from docx.shared import Pt, Inches

# PDF (simple text PDF)
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm

HEADER = "Medical Resubmission Letter"
ORG = "Andalusia Group"
FOOTER = "This resubmission includes medical justifications aligned with clinical best practices."

def _letters_header(doc: Document, title: str):
    h = doc.add_heading(title, level=1)
    for r in h.runs:
        r.font.size = Pt(16)

def _paragraph(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

def render_docx(df: pd.DataFrame, visit_id: str, patient_hint: str | None = None) -> bytes:
    """
    df columns expected: VisitServiceID, Service_Name, Reason, Justification, ContractorEnName, VisitStartDate
    """
    doc = Document()
    _letters_header(doc, HEADER)
    _paragraph(doc, ORG, bold=True)
    _paragraph(doc, f"Visit ID: {visit_id}")
    if patient_hint:
        _paragraph(doc, patient_hint)
    if not df.empty:
        vdate = str(df["VisitStartDate"].iloc[0]) if "VisitStartDate" in df.columns else "-"
        payer = str(df["ContractorEnName"].iloc[0]) if "ContractorEnName" in df.columns else "-"
        _paragraph(doc, f"Visit Date: {vdate}")
        _paragraph(doc, f"Payer: {payer}")
    doc.add_paragraph()  # spacer

    doc.add_heading("Justifications", level=2)
    for _, row in df.iterrows():
        sid = row.get("VisitServiceID", "")
        sname = row.get("Service_Name", "")
        reason = row.get("Reason", "")
        just = row.get("Justification", "")
        doc.add_paragraph(f"Service {sid}: {sname}", style=None).runs[0].bold = True
        _paragraph(doc, f"Rejection reason: {reason}")
        _paragraph(doc, just)
        doc.add_paragraph()

    doc.add_paragraph(FOOTER)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render_pdf(df: pd.DataFrame, visit_id: str, patient_hint: str | None = None) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    W, H = A4
    x, y = 20*mm, H - 20*mm
    line_h = 6*mm

    def write(line: str, bold=False):
        nonlocal y
        if y < 25*mm:
            c.showPage()
            y = H - 20*mm
        if bold:
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 11)
        c.drawString(x, y, line)
        y -= line_h

    write(HEADER, bold=True)
    write(ORG)
    write(f"Visit ID: {visit_id}")
    if patient_hint:
        write(patient_hint)
    if not df.empty:
        vdate = str(df["VisitStartDate"].iloc[0]) if "VisitStartDate" in df.columns else "-"
        payer = str(df["ContractorEnName"].iloc[0]) if "ContractorEnName" in df.columns else "-"
        write(f"Visit Date: {vdate}")
        write(f"Payer: {payer}")
    y -= 4*mm
    write("Justifications", bold=True)

    for _, row in df.iterrows():
        sid = row.get("VisitServiceID", "")
        sname = row.get("Service_Name", "")
        reason = row.get("Reason", "")
        just = row.get("Justification", "")
        write(f"Service {sid}: {sname}", bold=True)
        write(f"Rejection reason: {reason}")
        for chunk in wrap_text(just, 95):
            write(chunk)

        y -= 3*mm

    y -= 3*mm
    write(FOOTER)
    c.save()
    return buffer.getvalue()

def wrap_text(text: str, width: int) -> Iterable[str]:
    import textwrap
    return textwrap.wrap(text or "", width=width)